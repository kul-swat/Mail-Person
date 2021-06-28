from PyQt5 import QtCore, QtGui, QtWidgets
from docx.api import Document as document_api
from docx import Document as document
from docx.shared import Pt
import copy
import openpyxl
from docx2pdf import convert

class Ui_Init_Window(object):
    def __init__(self,wi,getpass,json,os,opx,sys):
        self.wi = wi
        self.getpass = getpass
        self.json = json
        self.os = os
        self.opx = opx
        self.sys = sys
    
    def wordAttachment(self):
        file = self.wi.CreateFileDialog(1,"All Files (*.*)|*.*|")
        file.DoModal()
        if self.lineEdit.text().strip():
            txtval = self.lineEdit.text()
            txtval = txtval + ',' + file.GetPathName()
            self.lineEdit.setText(txtval)
        else:
            self.lineEdit.setText(file.GetPathName())
    def mergeAttachment(self):
        file = self.wi.CreateFileDialog(1,"All Files (*.*)|*.*|")
        file.DoModal()
        if self.lineEdit_2.text().strip():
            txtval = self.lineEdit_2.text()
            txtval = txtval + ',' + file.GetPathName()
            self.lineEdit_2.setText(txtval)
        else:
            self.lineEdit_2.setText(file.GetPathName())
    def setupUi(self, MainWindow):
        MainWindow.setObjectName("MainWindow")
        MainWindow.resize(491, 335)
        self.centralwidget = QtWidgets.QWidget(MainWindow)
        self.centralwidget.setObjectName("centralwidget")
        self.label = QtWidgets.QLabel(self.centralwidget)
        self.label.setGeometry(QtCore.QRect(20, 20, 101, 31))
        self.label.setObjectName("label")
        self.label_2 = QtWidgets.QLabel(self.centralwidget)
        self.label_2.setGeometry(QtCore.QRect(51, 60, 61, 31))
        self.label_2.setObjectName("label_2")
        self.lineEdit = QtWidgets.QLineEdit(self.centralwidget)
        self.lineEdit.setGeometry(QtCore.QRect(130, 27, 261, 20))
        self.lineEdit.setObjectName("lineEdit")
        self.lineEdit_2 = QtWidgets.QLineEdit(self.centralwidget)
        self.lineEdit_2.setGeometry(QtCore.QRect(130, 66, 261, 20))
        self.lineEdit_2.setObjectName("lineEdit_2")
        self.lineEdit_2.setText(self.getpass.getuser())
        self.pushButton = QtWidgets.QPushButton(self.centralwidget)
        self.pushButton.setGeometry(QtCore.QRect(400, 25, 75, 23))
        self.pushButton.setObjectName("pushButton")
        self.pushButton.clicked.connect(self.wordAttachment)
        self.pushButton3 = QtWidgets.QPushButton(self.centralwidget)
        self.pushButton3.setGeometry(QtCore.QRect(400, 65, 75, 23))
        self.pushButton3.setObjectName("pushButton3")
        self.pushButton3.clicked.connect(self.mergeAttachment)
    
        self.pushButton_2 = QtWidgets.QPushButton(self.centralwidget)
        self.pushButton_2.setGeometry(QtCore.QRect(30, 170, 441, 31))
        self.pushButton_2.setObjectName("pushButton_2")
        self.pushButton_2.setStyleSheet("background-color: orange; font: bold 10pt")
        self.pushButton_2.clicked.connect(self.submit)
        
        
        MainWindow.setCentralWidget(self.centralwidget)
        self.menubar = QtWidgets.QMenuBar(MainWindow)
        self.menubar.setGeometry(QtCore.QRect(0, 0, 491, 21))
        self.menubar.setObjectName("menubar")
        MainWindow.setMenuBar(self.menubar)
        self.statusbar = QtWidgets.QStatusBar(MainWindow)
        self.statusbar.setObjectName("statusbar")
        MainWindow.setStatusBar(self.statusbar)

        self.retranslateUi(MainWindow)
        QtCore.QMetaObject.connectSlotsByName(MainWindow)
        self.default_val()      

    def default_val(self):
        self.lineEdit.setText("")
        self.lineEdit_2.setText("")      
                    
    def retranslateUi(self, MainWindow):
        _translate = QtCore.QCoreApplication.translate
        MainWindow.setWindowTitle(_translate("MainWindow", "MainWindow"))
        self.label.setText(_translate("MainWindow", "Word Template"))
        self.label_2.setText(_translate("MainWindow", "Merge File"))
        self.pushButton.setText(_translate("MainWindow", "Browse"))
        self.pushButton3.setText(_translate("MainWindow", "Browse"))
        
        
        self.pushButton_2.setText(_translate("MainWindow", "Submit"))


    def displayMsg(self,msgs,sub):
        msg = QtWidgets.QMessageBox()
        msg.setWindowTitle(sub)
        msg.setIcon(QtWidgets.QMessageBox.Information)
        msg.setText(msgs)
        msg.exec_()

    def modify_word(self,client_details):
        docs = copy.deepcopy(self.original)
        #for table
        for row in docs.tables[0].rows:
            for key in client_details.keys():
                if client_details[key]==None:
                    row.cells[0].text=row.cells[0].text.replace(f"«{key}»","")
                else:
                    row.cells[0].text=row.cells[0].text.replace(f"«{key}»",client_details[key])
                row.cells[0].paragraphs[0].style=self.style
            
        #for paragraph
        for line in docs.paragraphs:
            inline = line.runs
            for id,word in enumerate(inline):
                #print(key,id, word.text)
                for key in client_details.keys():
                    if key in word.text:
                        #print(key,id, word.text)
                        if client_details[key]==None:
                            word.text=word.text.replace(f"«{key}»","")
                        else:
                            word.text=word.text.replace(f"«{key}»",client_details[key])
        return docs

    def submit(self):
        state = dict()
        state['template_path']=self.lineEdit.text()
        state['excel_path']=self.lineEdit_2.text()
        
        self.original=document(state['template_path'])
        self.style = self.original.styles['Normal']
        self.font = self.style.font
        self.font.name = 'Arial'
        self.font.size = Pt(10)

        client_data={}
        wb=openpyxl.load_workbook(state['excel_path'])
        ws=wb.active
        max_rows = ws.max_row
        max_columns = ws.max_column

        for row in range(2,max_rows+1):
            if (ws.cell(row = row, column = 1)).value == None:
                break
            for column in range(1,max_columns+1):
                if (ws.cell(row = 1, column = column)).value==None:
                    break
                header_obj=ws.cell(row = 1, column = column)
                cell_obj=ws.cell(row = row, column = column)
                client_data[header_obj.value]=cell_obj.value
            modified_doc = self.modify_word(client_data)
            modified_doc.save(f'.\\temp\\{client_data["First_Name"]}.docx')
            convert(f'.\\temp\\{client_data["First_Name"]}.docx',f'.\\pdfs\\{client_data["First_Name"]}.pdf')
            #saving pdf path in hyperlink
            ws.cell(row=row, column=12).hyperlink = f'.\\pdfs\\{client_data["First_Name"]}.pdf'
            ws.cell(row=row, column=12).value = "pdf file"
            ws.cell(row=row, column=12).style = "Hyperlink"
            wb.save(state['excel_path'])

        self.displayMsg("All PDFs created successfully",'SUCCESS!!')

                


            